import os
from flask import Flask, request, render_template, send_file
from werkzeug.utils import secure_filename
from PIL import Image
import pytesseract
from pdf2image import convert_from_path
from docx import Document

app = Flask(__name__)
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/convert', methods=['POST'])
def convert():
    file = request.files.get('file')
    if not file:
        return "No file uploaded", 400

    filename = secure_filename(file.filename)
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(filepath)

    doc = Document()

    try:
        if filename.lower().endswith('.pdf'):
            images = convert_from_path(filepath)
            for img in images:
                text = pytesseract.image_to_string(img)
                doc.add_paragraph(text)
        else:
            img = Image.open(filepath)
            text = pytesseract.image_to_string(img)
            doc.add_paragraph(text)
    except Exception as e:
        return f"Error during processing: {e}", 500

    output_path = os.path.join(app.config['UPLOAD_FOLDER'], 'output.docx')
    doc.save(output_path)
    return send_file(output_path, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)
